import os
import tempfile
import unittest

from app.core.models import FileItem
from core.workers.merge.merge_mixin import MergeMixin


class _SignalStub:
    def __init__(self):
        self.emitted = []

    def emit(self, value):
        self.emitted.append(value)


class _DummyMergeWorker(MergeMixin):
    def __init__(self):
        self.files = []
        self.errors = []
        self.merge_output_format = "docx"
        self.merge_output_path = ""
        self.status = _SignalStub()
        self.progress = _SignalStub()
        self.finished = _SignalStub()
        self.error = _SignalStub()
        self._cancel_requested = False

    def _should_cancel(self):
        return self._cancel_requested

    def _record_error(self, file_item, message):
        self.errors.append({"file": file_item, "message": message})

    def _get_unique_path(self, path):
        if not os.path.exists(path):
            return path
        base, ext = os.path.splitext(path)
        return f"{base}_1{ext}"


class MergeMixinTests(unittest.TestCase):
    def test_merge_docx_files(self):
        from docx import Document

        worker = _DummyMergeWorker()
        with tempfile.TemporaryDirectory() as tmpdir:
            first = os.path.join(tmpdir, "first.docx")
            second = os.path.join(tmpdir, "second.docx")

            doc = Document()
            doc.add_paragraph("First")
            doc.save(first)

            doc = Document()
            doc.add_paragraph("Second")
            doc.save(second)

            worker.files = [FileItem(first), FileItem(second)]
            output = worker._merge_files_to_target()

            self.assertTrue(os.path.exists(output))
            merged = Document(output)
            text = "\n".join(paragraph.text for paragraph in merged.paragraphs)
            self.assertIn("First", text)
            self.assertIn("Second", text)

    def test_merge_uses_requested_output_path(self):
        from docx import Document

        worker = _DummyMergeWorker()
        with tempfile.TemporaryDirectory() as tmpdir:
            first = os.path.join(tmpdir, "first.docx")
            second = os.path.join(tmpdir, "second.docx")
            requested = os.path.join(tmpdir, "custom-name.docx")

            for path, value in ((first, "First"), (second, "Second")):
                doc = Document()
                doc.add_paragraph(value)
                doc.save(path)

            worker.files = [FileItem(first), FileItem(second)]
            worker.merge_output_path = requested

            self.assertEqual(worker._merge_files_to_target(), requested)
            self.assertTrue(os.path.exists(requested))


if __name__ == "__main__":
    unittest.main()
