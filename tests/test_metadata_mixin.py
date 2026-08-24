import os
import tempfile
import unittest
import zipfile
from xml.etree import ElementTree as ET

from core.workers.metadata.metadata_mixin import MetadataMixin


class _MetadataHarness(MetadataMixin):
    pass


class MetadataMixinTests(unittest.TestCase):
    def setUp(self):
        self.worker = _MetadataHarness()

    def test_pdf_selective_cleanup_preserves_unselected_fields(self):
        import fitz

        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "sample.pdf")
            document = fitz.open()
            document.new_page().insert_text((72, 72), "Visible content")
            document.set_metadata(
                {
                    "title": "Keep title",
                    "author": "Remove author",
                    "subject": "Keep subject",
                    "keywords": "Keep keywords",
                }
            )
            document.save(path)
            document.close()

            self.worker._remove_pdf_metadata(path, False, {"author"})

            result = fitz.open(path)
            try:
                metadata = result.metadata
                self.assertEqual(metadata.get("author"), "")
                self.assertEqual(metadata.get("title"), "Keep title")
                self.assertEqual(metadata.get("subject"), "Keep subject")
                self.assertIn("Visible content", result[0].get_text())
            finally:
                result.close()

    def test_docx_selective_cleanup_preserves_content_and_title(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "sample.docx")
            document = Document()
            document.add_paragraph("Document body")
            document.core_properties.author = "Remove author"
            document.core_properties.title = "Keep title"
            document.core_properties.keywords = "Keep keywords"
            document.save(path)

            self.worker._remove_docx_metadata(path, False, {"author"})

            result = Document(path)
            self.assertEqual(result.core_properties.author, "")
            self.assertEqual(result.core_properties.title, "Keep title")
            self.assertEqual(result.core_properties.keywords, "Keep keywords")
            self.assertEqual(result.paragraphs[0].text, "Document body")

    def test_docx_full_cleanup_removes_property_parts(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "sample.docx")
            document = Document()
            document.add_paragraph("Document body")
            document.core_properties.author = "Author"
            document.core_properties.title = "Title"
            document.save(path)

            self.worker._remove_docx_metadata(path, True, set())

            with zipfile.ZipFile(path, "r") as archive:
                names = set(archive.namelist())
                self.assertNotIn("docProps/core.xml", names)
                self.assertNotIn("docProps/app.xml", names)
                self.assertNotIn("docProps/custom.xml", names)

            result = Document(path)
            self.assertEqual(result.paragraphs[0].text, "Document body")

    def test_odt_selective_cleanup_removes_only_selected_elements(self):
        metadata_xml = b'''<?xml version="1.0" encoding="UTF-8"?>
<office:document-meta
 xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
 xmlns:dc="http://purl.org/dc/elements/1.1/"
 xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0">
 <office:meta>
  <dc:title>Keep title</dc:title>
  <dc:creator>Remove author</dc:creator>
  <meta:keyword>Keep keyword</meta:keyword>
 </office:meta>
</office:document-meta>'''

        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "sample.odt")
            with zipfile.ZipFile(path, "w") as archive:
                archive.writestr(
                    "mimetype",
                    "application/vnd.oasis.opendocument.text",
                    compress_type=zipfile.ZIP_STORED,
                )
                archive.writestr("meta.xml", metadata_xml)

            self.worker._remove_odt_metadata(path, False, {"author"})

            with zipfile.ZipFile(path, "r") as archive:
                root = ET.fromstring(archive.read("meta.xml"))
            local_names = [element.tag.rsplit("}", 1)[-1] for element in root.iter()]
            self.assertNotIn("creator", local_names)
            self.assertIn("title", local_names)
            self.assertIn("keyword", local_names)


if __name__ == "__main__":
    unittest.main()
