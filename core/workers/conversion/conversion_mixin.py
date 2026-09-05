import html
import os
import re
import tempfile
import textwrap
import threading
import time
import urllib.parse
from collections.abc import Callable
from concurrent.futures import ThreadPoolExecutor, as_completed

from app.core.app_utils import _debug_log
from app.core.deps import (
    HAS_ODF_PYTHON,
    HAS_PDF_TO_WORD,
    HAS_PIL,
    HAS_PYMUPDF,
    HAS_WORD_TO_PDF,
    Image,
    OpenDocumentText,
    load,
    pdf2docx,
    teletype,
    text,
)
from app.core.models import FileItem
from app.core.conversion_formats import (
    DOCUMENT_CATEGORY,
    IMAGE_CATEGORY,
    format_for_path,
    source_formats_for_category,
    suffix_for_format,
)

from .image_encoding_mixin import ImageEncodingMixin

_TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "cp1251", "cp1252", "latin-1")

_HTML_BLOCK_TAGS = frozenset({"p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6"})

_HTML_END_BLOCK_TAGS = _HTML_BLOCK_TAGS - {"br"}


_WORD_WARMUP_LOCK = threading.Lock()
_CONVERSION_OUTPUT_PATH_LOCK = threading.Lock()
_WORD_WARMUP_DONE = False
_WD_EXPORT_FORMAT_PDF = 17


_WORD_DIRECT_SOURCE_FORMATS = frozenset({"DOC", "DOCX", "ODT", "RTF", "TXT", "HTML"})
_WORD_DIRECT_TARGET_FORMATS = frozenset({"DOC", "DOCX", "PDF", "ODT", "RTF", "HTML"})
_WORD_SAVE_FORMATS: dict[str, int] = {
    "DOC": 0,
    "DOCX": 16,
    "RTF": 6,
    "HTML": 10,
    "ODT": 23,
}


def _close_word_document(document) -> None:
    if document is None:
        return
    try:
        document.Close(False)
    except Exception as error:
        _debug_log(f"Не удалось закрыть документ Word: {error}")


def _quit_word_application(word_application) -> None:
    if word_application is None:
        return
    try:
        word_application.Quit()
    except Exception as error:
        _debug_log(f"Не удалось закрыть Microsoft Word: {error}")


def _uninitialize_com(pythoncom_module) -> None:
    try:
        pythoncom_module.CoUninitialize()
    except Exception as error:
        _debug_log(f"Не удалось завершить COM-сеанс Word: {error}")


def prewarm_word_background(status_callback=None, log_callback=None) -> bool:
    """Запускает фоновую подготовку Microsoft Word один раз за сессию."""
    global _WORD_WARMUP_DONE
    if os.name != "nt" or not HAS_WORD_TO_PDF:
        return False
    if _WORD_WARMUP_DONE:
        return False

    with _WORD_WARMUP_LOCK:
        if _WORD_WARMUP_DONE:
            return False
        try:
            if callable(status_callback):
                status_callback("Фоновая подготовка Microsoft Word...")
            import pythoncom
            import win32com.client as win32

            pythoncom.CoInitialize()
            word_app = None
            try:
                word_app = win32.DispatchEx("Word.Application")
                word_app.Visible = False
                word_app.DisplayAlerts = 0
                _WORD_WARMUP_DONE = True
                return True
            finally:
                _quit_word_application(word_app)
                _uninitialize_com(pythoncom)
        except Exception as e:
            msg = f"Не удалось запустить фоновую подготовку Microsoft Word: {e}"
            if callable(log_callback):
                try:
                    log_callback(msg)
                except Exception:
                    _debug_log(msg)
            else:
                _debug_log(msg)
            return False


class ConversionMixin(ImageEncodingMixin):
    CONVERTED_FOLDER_NAME = "Конвертированные"

    def _conversion_output_directory(self, file: FileItem) -> str:
        custom_dir = str(getattr(self, "conversion_output_dir", "") or "").strip()
        if custom_dir:
            output_dir = os.path.abspath(os.path.expanduser(custom_dir))
        else:
            source_dir = os.path.dirname(os.path.abspath(file.path))
            output_dir = os.path.join(source_dir, self.CONVERTED_FOLDER_NAME)
        os.makedirs(output_dir, exist_ok=True)
        return output_dir

    def _conversion_output_path(
        self,
        file: FileItem,
        target_ext: str,
        *,
        reference_file: FileItem | None = None,
    ) -> str:
        target_ext = str(target_ext or "").strip()
        if target_ext and not target_ext.startswith("."):
            target_ext = f".{target_ext}"
        reference = reference_file or file
        stem = os.path.splitext(os.path.basename(reference.path))[0]
        base_path = os.path.join(
            self._conversion_output_directory(reference),
            f"{stem}{target_ext}",
        )

        with _CONVERSION_OUTPUT_PATH_LOCK:
            reserved = getattr(self, "_conversion_reserved_paths", None)
            if reserved is None:
                reserved = set()
                self._conversion_reserved_paths = reserved

            candidate = base_path
            base, ext = os.path.splitext(base_path)
            counter = 1
            while os.path.exists(candidate) or os.path.normcase(candidate) in reserved:
                candidate = f"{base}_{counter}{ext}"
                counter += 1
            reserved.add(os.path.normcase(candidate))
        return candidate

    def _release_conversion_output_path(self, path: str) -> None:
        reserved = getattr(self, "_conversion_reserved_paths", None)
        if reserved is not None and path:
            reserved.discard(os.path.normcase(path))

    def _discard_conversion_output(self, path: str) -> None:
        """Удаляет частичный результат и освобождает зарезервированное имя."""
        if not path:
            return
        try:
            if os.path.exists(path):
                os.remove(path)
        except OSError as error:
            _debug_log(f"Не удалось удалить частичный результат {path}: {error}")
        finally:
            self._release_conversion_output_path(path)

    def _report_file_conversion_error(
        self,
        file: FileItem,
        error: Exception,
    ) -> None:
        """Сохраняет и публикует ошибку конвертации одного файла."""
        message = f"Ошибка конвертации файла {file.name}: {error}"
        self._record_error(file, message)
        self.error.emit(message)

    @staticmethod
    def _converted_file_item(path: str | None) -> FileItem | None:
        if path and os.path.exists(path):
            return FileItem(path)
        return None

    def _get_conversion_handler(self) -> Callable[[FileItem], str | None] | None:
        handlers: dict[str, Callable[[FileItem], str | None]] = {
            "word_to_odt": self._convert_word_to_odt,
            "odt_to_word": self._convert_odt_to_word,
            "odt_to_pdf": self._convert_odt_to_pdf,
            "pdf_to_odt": self._convert_pdf_to_odt,
            "pdf_to_image": self._convert_pdf_to_image,
            "pdf_to_images": self._convert_pdf_to_image,
            "image_to_pdf": self._convert_image_to_pdf,
            "image_to_image": lambda file: self._convert_image_to_image(
                file, self.conversion_format
            ),
            "auto_image": lambda file: self._convert_image_auto(file, self.conversion_format),
            "auto_document": lambda file: self._convert_document_auto(
                file, self.conversion_format
            ),
        }
        return handlers.get(self.conversion_type)

    @staticmethod
    def _normalize_word_com_path(path: str) -> str:
        value = str(path or "").strip().strip('"')
        if not value:
            return value
        if value.lower().startswith("file:///"):
            value = value[8:]
        value = urllib.parse.unquote(value)
        if value.startswith("/") and len(value) >= 3 and value[2] == ":":
            value = value[1:]
        value = value.replace("/", "\\")
        return os.path.normpath(value)

    def _convert_image_to_image(self, file: FileItem, target_format: str) -> str:
        if not HAS_PIL:
            raise Exception("Установите Pillow для конвертации изображений")

        source_format = format_for_path(file.path)
        if source_format not in source_formats_for_category(IMAGE_CATEGORY):
            extension = os.path.splitext(file.path)[1]
            raise Exception(f"Неподдерживаемый формат изображения: {extension}")
        target_format = str(target_format or "").upper()
        target_ext = suffix_for_format(target_format)
        if not target_ext:
            raise Exception(f"Неизвестный формат изображения: {target_format}")
        if source_format == target_format:
            return None
        if target_format == "PDF":
            return self._convert_image_to_pdf(file)

        output_path = self._conversion_output_path(file, target_ext)
        image = None
        try:
            if source_format == "SVG":
                image = self._load_svg_as_pillow_image(file.path)
            else:
                image = Image.open(file.path)
            self._save_pillow_image(image, output_path, target_format)
            return output_path
        except Exception as error:
            self._discard_conversion_output(output_path)
            raise Exception(
                f"Ошибка конвертации изображения {source_format} → {target_format}: {error}"
            ) from error
        finally:
            if image is not None:
                try:
                    image.close()
                except Exception as close_error:
                    _debug_log(f"Не удалось закрыть изображение {file.path}: {close_error}")

    def _convert_image_auto(self, file: FileItem, target_format: str) -> str:
        source_format = format_for_path(file.path)
        if source_format not in source_formats_for_category(IMAGE_CATEGORY):
            raise Exception("Файл не является поддерживаемым изображением")
        if source_format == str(target_format or "").upper():
            self.status.emit(f"Пропущен {file.name}: уже {target_format}")
            return None
        return self._convert_image_to_image(file, target_format)

    @staticmethod
    def _read_text_file(path: str) -> str:
        with open(path, "rb") as stream:
            data = stream.read()
        for encoding in _TEXT_ENCODINGS:
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")

    @staticmethod
    def _html_to_text(value: str) -> str:
        from html.parser import HTMLParser

        class _Extractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.parts = []

            def handle_starttag(self, tag, attrs):
                if tag.lower() in _HTML_BLOCK_TAGS:
                    self.parts.append("\n")

            def handle_endtag(self, tag):
                if tag.lower() in _HTML_END_BLOCK_TAGS:
                    self.parts.append("\n")

            def handle_data(self, data):
                self.parts.append(data)

        parser = _Extractor()
        parser.feed(value)
        text_value = "".join(parser.parts)
        text_value = re.sub(r"[ \t]+\n", "\n", text_value)
        text_value = re.sub(r"\n{3,}", "\n\n", text_value)
        return html.unescape(text_value).strip()

    @staticmethod
    def _rtf_to_text(value: str) -> str:
        # Достаточный резервный разбор для обычного RTF. При наличии Word используется
        # COM-конвертация, которая сохраняет форматирование намного точнее.
        value = re.sub(
            r"\\'([0-9a-fA-F]{2})",
            lambda match: bytes([int(match.group(1), 16)]).decode("cp1251", errors="replace"),
            value,
        )
        value = re.sub(
            r"\\u(-?\d+)\??",
            lambda match: chr(int(match.group(1)) % 65536),
            value,
        )
        value = value.replace(r"\par", "\n").replace(r"\line", "\n").replace(r"\tab", "\t")
        value = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", value)
        value = value.replace(r"\{", "{").replace(r"\}", "}").replace(r"\\", "\\")
        value = value.replace("{", "").replace("}", "")
        value = re.sub(r"\n{3,}", "\n\n", value)
        return value.strip()

    def _extract_text_via_word(self, path: str) -> str:
        if os.name != "nt" or not HAS_WORD_TO_PDF:
            raise Exception("Для чтения DOC требуется Microsoft Word и pywin32")
        import pythoncom
        import win32com.client as win32

        pythoncom.CoInitialize()
        word_app = None
        document = None
        try:
            word_app = win32.DispatchEx("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = 0
            document = word_app.Documents.Open(
                self._normalize_word_com_path(path),
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                Revert=False,
                Visible=False,
                OpenAndRepair=True,
            )
            return str(document.Content.Text or "")
        finally:
            _close_word_document(document)
            _quit_word_application(word_app)
            _uninitialize_com(pythoncom)

    def _extract_document_text(self, file: FileItem) -> str:
        source_format = format_for_path(file.path)
        if source_format in {"TXT", "MD"}:
            return self._read_text_file(file.path)
        if source_format == "HTML":
            return self._html_to_text(self._read_text_file(file.path))
        if source_format == "RTF":
            try:
                if os.name == "nt" and HAS_WORD_TO_PDF:
                    return self._extract_text_via_word(file.path)
            except Exception as error:
                _debug_log(f"Word не смог прочитать RTF, используется fallback: {error}")
            return self._rtf_to_text(self._read_text_file(file.path))
        if source_format == "DOC":
            return self._extract_text_via_word(file.path)
        if source_format == "DOCX":
            try:
                import docx

                document = docx.Document(file.path)
                parts = [paragraph.text for paragraph in document.paragraphs]
                for table in document.tables:
                    for row in table.rows:
                        parts.append("\t".join(cell.text for cell in row.cells))
                return "\n".join(parts)
            except Exception as error:
                raise Exception(f"Ошибка чтения DOCX: {error}") from error
        if source_format == "ODT":
            if not HAS_ODF_PYTHON:
                raise Exception("Установите odfpy для чтения ODT")
            try:
                odt_doc = load(file.path)
                paragraphs: list[str] = []
                for element in odt_doc.getElementsByType(text.P):
                    paragraph = teletype.extractText(element)
                    if paragraph.strip():
                        paragraphs.append(paragraph)
                return "\n".join(paragraphs)
            except Exception as error:
                raise Exception(f"Ошибка чтения ODT: {error}") from error
        if source_format in {"PDF", "EPUB", "FB2", "XPS", "MOBI"}:
            if not HAS_PYMUPDF:
                raise Exception("Установите PyMuPDF для чтения этого документа")
            try:
                import pymupdf as fitz

                with fitz.open(file.path) as document:
                    return "\n".join(page.get_text("text") for page in document)
            except Exception as error:
                raise Exception(f"Ошибка извлечения текста из {source_format}: {error}") from error
        raise Exception(f"Неподдерживаемый формат документа: {source_format or file.path}")

    @staticmethod
    def _write_plain_text_file(output_path: str, content: str) -> None:
        with open(output_path, "w", encoding="utf-8", newline="") as stream:
            stream.write(content)

    @staticmethod
    def _write_html_file(output_path: str, content: str) -> None:
        paragraphs: list[str] = []
        for paragraph in re.split(r"\n\s*\n", content):
            escaped = html.escape(paragraph.strip()).replace("\n", "<br>\n")
            if escaped:
                paragraphs.append(f"<p>{escaped}</p>")

        with open(output_path, "w", encoding="utf-8", newline="") as stream:
            stream.write(
                '<!doctype html><html><head><meta charset="utf-8"></head><body>\n'
            )
            stream.write("\n".join(paragraphs))
            stream.write("\n</body></html>")

    def _write_rtf_file(self, output_path: str, content: str) -> None:
        escaped = "".join(self._rtf_escape_char(char) for char in content)
        escaped = escaped.replace("\n", r"\par" + "\n")
        with open(
            output_path,
            "w",
            encoding="ascii",
            errors="ignore",
            newline="",
        ) as stream:
            stream.write(r"{\rtf1\ansi\deff0\uc1 ")
            stream.write(escaped)
            stream.write("}")

    @staticmethod
    def _write_docx_file(output_path: str, content: str) -> None:
        try:
            import docx

            document = docx.Document()
            for paragraph in content.splitlines() or [""]:
                document.add_paragraph(paragraph)
            document.save(output_path)
        except Exception as error:
            raise Exception(f"Ошибка создания DOCX: {error}") from error

    @staticmethod
    def _write_odt_file(output_path: str, content: str) -> None:
        if not HAS_ODF_PYTHON:
            raise Exception("Установите odfpy для создания ODT")
        try:
            document = OpenDocumentText()
            for paragraph in content.splitlines() or [""]:
                document.text.addElement(text.P(text=paragraph))
            document.save(output_path)
        except Exception as error:
            raise Exception(f"Ошибка создания ODT: {error}") from error

    def _write_text_output(
        self,
        file: FileItem,
        target_format: str,
        content: str,
    ) -> str:
        normalized_target = str(target_format or "").upper()
        if normalized_target == "PDF":
            # PDF создаёт собственный путь результата. Ранний выход исключает
            # двойное резервирование имени и появление лишнего суффикса ``_1``.
            return self._write_text_pdf(file, content)

        target_ext = suffix_for_format(normalized_target)
        if not target_ext:
            raise Exception(f"Нельзя сохранить текст в {normalized_target}")

        output_path = self._conversion_output_path(file, target_ext)
        writers: dict[str, Callable[[str, str], None]] = {
            "TXT": self._write_plain_text_file,
            "MD": self._write_plain_text_file,
            "HTML": self._write_html_file,
            "RTF": self._write_rtf_file,
            "DOCX": self._write_docx_file,
            "ODT": self._write_odt_file,
        }
        writer = writers.get(normalized_target)
        if writer is None:
            self._release_conversion_output_path(output_path)
            raise Exception(f"Нельзя сохранить текст в {normalized_target}")

        try:
            writer(output_path, content)
            return output_path
        except Exception:
            self._discard_conversion_output(output_path)
            raise

    @staticmethod
    def _rtf_escape_char(char: str) -> str:
        """Экранирует символ для простого RTF fallback-экспортера."""
        if char in "\\{}":
            return "\\" + char
        codepoint = ord(char)
        if codepoint <= 127:
            return char
        signed = codepoint if codepoint < 32768 else codepoint - 65536
        return f"\\u{signed}?"

    @staticmethod
    def _find_unicode_font() -> str | None:
        candidates = [
            r"C:\Windows\Fonts\segoeui.ttf",
            r"C:\Windows\Fonts\arial.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
        ]
        for candidate in candidates:
            if os.path.exists(candidate):
                return candidate
        return None

    def _write_text_pdf(self, file: FileItem, content: str) -> str:
        if not HAS_PYMUPDF:
            raise Exception("Установите PyMuPDF для создания PDF")

        import pymupdf as fitz

        output_path = self._conversion_output_path(file, ".pdf")
        document = fitz.open()
        try:
            font_path = self._find_unicode_font()
            font_kwargs = (
                {"fontname": "multifora", "fontfile": font_path}
                if font_path
                else {"fontname": "helv"}
            )
            page = None
            y_position = 0.0
            page_width, page_height = 595.0, 842.0
            margin = 48.0
            line_height = 15.0
            max_chars = 92

            def new_page() -> None:
                nonlocal page, y_position
                page = document.new_page(width=page_width, height=page_height)
                y_position = margin

            new_page()
            normalized_content = content.replace("\r\n", "\n").replace("\r", "\n")
            for paragraph in normalized_content.split("\n"):
                lines = textwrap.wrap(
                    paragraph,
                    width=max_chars,
                    replace_whitespace=False,
                    drop_whitespace=False,
                ) or [""]
                for line in lines:
                    if y_position > page_height - margin:
                        new_page()
                    page.insert_text(
                        (margin, y_position),
                        line,
                        fontsize=10.5,
                        **font_kwargs,
                    )
                    y_position += line_height
            document.save(output_path)
            return output_path
        except Exception:
            self._discard_conversion_output(output_path)
            raise
        finally:
            document.close()

    def _try_word_direct_conversion(
        self,
        file: FileItem,
        target_format: str,
        *,
        output_reference: FileItem | None = None,
    ) -> str | None:
        source_format = format_for_path(file.path)
        target_format = str(target_format or "").upper()
        if os.name != "nt" or not HAS_WORD_TO_PDF:
            return None
        if source_format not in _WORD_DIRECT_SOURCE_FORMATS:
            return None
        if target_format not in _WORD_DIRECT_TARGET_FORMATS:
            return None
        target_ext = suffix_for_format(target_format)
        output_path = self._conversion_output_path(
            file,
            target_ext,
            reference_file=output_reference,
        )

        import pythoncom
        import win32com.client as win32

        pythoncom.CoInitialize()
        word_app = None
        document = None
        try:
            word_app = win32.DispatchEx("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = 0
            document = word_app.Documents.Open(
                self._normalize_word_com_path(file.path),
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                Revert=False,
                Visible=False,
                OpenAndRepair=True,
            )
            normalized_dst = self._normalize_word_com_path(output_path)
            if target_format == "PDF":
                document.ExportAsFixedFormat(
                    OutputFileName=normalized_dst,
                    ExportFormat=_WD_EXPORT_FORMAT_PDF,
                    OpenAfterExport=False,
                )
            else:
                document.SaveAs2(
                    FileName=normalized_dst,
                    FileFormat=_WORD_SAVE_FORMATS[target_format],
                )
            if not os.path.exists(normalized_dst):
                raise Exception("Microsoft Word не создал выходной файл")
            return output_path
        except Exception as error:
            _debug_log(
                f"Прямая конвертация Word {source_format} → {target_format} не удалась: {error}"
            )
            self._discard_conversion_output(output_path)
            return None
        finally:
            _close_word_document(document)
            _quit_word_application(word_app)
            _uninitialize_com(pythoncom)

    def _convert_pymupdf_document_to_pdf(self, file: FileItem) -> str | None:
        source_format = format_for_path(file.path)
        if source_format not in {"EPUB", "FB2", "XPS", "MOBI"} or not HAS_PYMUPDF:
            return None
        try:
            import pymupdf as fitz

            output_path = self._conversion_output_path(file, ".pdf")
            with fitz.open(file.path) as document:
                pdf_bytes = document.convert_to_pdf()
            with open(output_path, "wb") as stream:
                stream.write(pdf_bytes)
            return output_path
        except Exception as error:
            if "output_path" in locals():
                self._discard_conversion_output(output_path)
            _debug_log(
                f"PyMuPDF не смог напрямую создать PDF из {source_format}: {error}"
            )
            return None

    def _convert_document_to_doc(
        self,
        file: FileItem,
        source_format: str,
    ) -> str:
        direct_result = self._try_word_direct_conversion(file, "DOC")
        if direct_result:
            return direct_result
        if os.name != "nt" or not HAS_WORD_TO_PDF:
            raise Exception("Для сохранения в DOC требуется Microsoft Word и pywin32")

        with tempfile.TemporaryDirectory(prefix="multifora_doc_") as temp_dir:
            generated_intermediate = source_format == "PDF"
            if generated_intermediate:
                temp_docx = self._convert_pdf_to_word(file)
                if not temp_docx:
                    raise Exception("Не удалось подготовить промежуточный DOCX")
            else:
                temp_docx = os.path.join(temp_dir, "intermediate.docx")
                content = self._extract_document_text(file)
                self._write_docx_file(temp_docx, content)

            try:
                result = self._try_word_direct_conversion(
                    FileItem(temp_docx),
                    "DOC",
                    output_reference=file,
                )
                if not result:
                    raise Exception("Microsoft Word не смог создать DOC")
                return result
            finally:
                if generated_intermediate:
                    try:
                        os.remove(temp_docx)
                    except FileNotFoundError:
                        pass
                    except OSError as error:
                        _debug_log(
                            f"Не удалось удалить промежуточный DOCX {temp_docx}: {error}"
                        )
                    finally:
                        self._release_conversion_output_path(temp_docx)

    def _convert_document_to_pdf(
        self,
        file: FileItem,
        source_format: str,
    ) -> str:
        direct_result = self._try_word_direct_conversion(file, "PDF")
        if direct_result:
            return direct_result

        pymupdf_result = self._convert_pymupdf_document_to_pdf(file)
        if pymupdf_result:
            return pymupdf_result

        if source_format == "DOC" and (os.name != "nt" or not HAS_WORD_TO_PDF):
            raise Exception("Для DOC → PDF требуется Microsoft Word и pywin32")

        content = self._extract_document_text(file)
        return self._write_text_pdf(file, content)

    def _convert_document_auto(
        self,
        file: FileItem,
        target_format: str,
    ) -> str | None:
        source_format = format_for_path(file.path)
        normalized_target = str(target_format or "").upper()
        if source_format not in source_formats_for_category(DOCUMENT_CATEGORY):
            raise Exception("Файл не является поддерживаемым документом")
        if source_format == normalized_target:
            self.status.emit(f"Пропущен {file.name}: уже {normalized_target}")
            return None

        # Для PDF сохраняем специализированные конвертеры, потому что они лучше
        # восстанавливают структуру документа, чем текстовый резервный маршрут.
        if source_format == "PDF" and normalized_target == "DOCX":
            return self._convert_pdf_to_word(file)
        if source_format == "PDF" and normalized_target == "ODT":
            return self._convert_pdf_to_odt(file)
        if normalized_target == "DOC":
            return self._convert_document_to_doc(file, source_format)
        if normalized_target == "PDF":
            return self._convert_document_to_pdf(file, source_format)

        # Word используется первым для совместимых форматов, чтобы максимально
        # сохранить таблицы, шрифты и исходное форматирование.
        if normalized_target in {"DOCX", "ODT", "RTF", "HTML"}:
            direct_result = self._try_word_direct_conversion(file, normalized_target)
            if direct_result:
                return direct_result

        content = self._extract_document_text(file)
        return self._write_text_output(file, normalized_target, content)

    def _convert_files(self) -> None:
        total = len(self.files)
        results: list[FileItem] = []

        # Для тяжёлых сценариев отдельные пакетные обработчики переиспользуют
        # один Word-процесс или ограниченный пул потоков.
        if self.conversion_type == "word_to_pdf":
            self._emit_finished(self._convert_word_to_pdf_batch(), [])
            return
        if self.conversion_type == "pdf_to_word":
            self._emit_finished(self._convert_pdf_to_word_batch(), [])
            return

        handler = self._get_conversion_handler()
        for index, file in enumerate(self.files):
            if self._should_cancel():
                self.status.emit("Операция отменена пользователем")
                self._emit_finished(results, [])
                return
            if not file.is_file:
                continue

            self.status.emit(f"Конвертация: {file.name}")
            try:
                if handler is None:
                    raise Exception(
                        f"Неизвестная операция конвертации: {self.conversion_type}"
                    )
                converted_item = self._converted_file_item(handler(file))
                if converted_item is not None:
                    results.append(converted_item)
            except Exception as error:
                self._report_file_conversion_error(file, error)

            self.progress.emit(int((index + 1) / total * 100))

        self._emit_finished(results, [])

    def _convert_sequential_batch(
        self,
        files: list[FileItem],
        converter: Callable[[FileItem], str | None],
    ) -> list[FileItem]:
        """Выполняет однотипную конвертацию последовательно с общим UI-циклом."""
        total = len(files)
        results: list[FileItem] = []
        for index, file in enumerate(files):
            if self._should_cancel():
                self.status.emit("Операция отменена пользователем")
                break
            if not file.is_file:
                continue

            self.status.emit(f"Конвертация: {file.name}")
            try:
                converted_item = self._converted_file_item(converter(file))
                if converted_item is not None:
                    results.append(converted_item)
            except Exception as error:
                self._report_file_conversion_error(file, error)
            self.progress.emit(int((index + 1) / total * 100))
        return results

    def _convert_word_file_with_application(self, word_app, file: FileItem) -> str:
        """Экспортирует один DOC/DOCX через уже запущенный экземпляр Word."""
        if not file.path.lower().endswith((".doc", ".docx")):
            raise Exception("Поддерживаются только DOC/DOCX")

        pdf_path = self._conversion_output_path(file, ".pdf")
        try:
            normalized_src = self._normalize_word_com_path(file.path)
            normalized_dst = self._normalize_word_com_path(pdf_path)
            if not os.path.exists(normalized_src):
                raise FileNotFoundError(f"Файл не найден: {normalized_src}")

            document = None
            try:
                document = word_app.Documents.Open(
                    normalized_src,
                    ConfirmConversions=False,
                    ReadOnly=True,
                    AddToRecentFiles=False,
                    Revert=False,
                    Visible=False,
                    OpenAndRepair=True,
                )
                document.ExportAsFixedFormat(
                    OutputFileName=normalized_dst,
                    ExportFormat=_WD_EXPORT_FORMAT_PDF,
                    OpenAfterExport=False,
                )
            finally:
                _close_word_document(document)

            if not os.path.exists(normalized_dst):
                raise Exception("Word не создал выходной PDF-файл")
            return pdf_path
        except Exception:
            self._discard_conversion_output(pdf_path)
            raise

    def _convert_word_to_pdf_batch(self) -> list[FileItem]:
        """Быстрый пакетный DOC/DOCX -> PDF через один скрытый экземпляр Word."""
        total = len(self.files)
        results: list[FileItem] = []
        if total == 0:
            return results

        if os.name != "nt":
            return self._convert_sequential_batch(
                list(self.files),
                self._convert_word_to_pdf,
            )

        try:
            import pythoncom
            import win32com.client as win32
        except (ImportError, OSError) as error:
            msg = f"Скрытая конвертация Word недоступна (нужен pywin32): {error}"
            self._record_error(None, msg)
            self.error.emit(msg)
            return results

        self._warmup_word()
        word_app = None
        pythoncom.CoInitialize()
        try:
            # Один COM-экземпляр Word на всю пачку файлов заметно ускоряет конвертацию.
            word_app = win32.DispatchEx("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = 0

            for index, file in enumerate(self.files):
                if self._should_cancel():
                    self.status.emit("Операция отменена пользователем")
                    break
                if not file.is_file:
                    continue

                self.status.emit(f"Конвертация: {file.name}")
                try:
                    pdf_path = self._convert_word_file_with_application(word_app, file)
                    results.append(FileItem(pdf_path))
                except Exception as error:
                    self._report_file_conversion_error(file, error)
                self.progress.emit(int((index + 1) / total * 100))
        finally:
            _quit_word_application(word_app)
            _uninitialize_com(pythoncom)

        return results

    def _convert_pdf_to_word_batch(self) -> list[FileItem]:
        files = [f for f in self.files if getattr(f, "is_file", False)]
        total = len(files)
        results: list[FileItem] = []
        if total == 0:
            return results

        if not HAS_PDF_TO_WORD:
            msg = "Установите pdf2docx"
            self._record_error(None, msg)
            self.error.emit(msg)
            return results

        max_workers = min(4, max(1, ((os.cpu_count() or 2) // 2)))
        if total <= 1 or max_workers == 1:
            return self._convert_sequential_batch(files, self._convert_pdf_to_word)

        processed = 0
        with ThreadPoolExecutor(
            max_workers=max_workers,
            thread_name_prefix="pdf2docx",
        ) as executor:
            future_map = {}
            for file in files:
                if file.path.lower().endswith(".pdf"):
                    future_map[executor.submit(self._convert_pdf_to_word, file)] = file
                else:
                    processed += 1
                    self.progress.emit(int(processed / total * 100))

            for future in as_completed(future_map):
                file = future_map[future]
                if self._should_cancel():
                    for pending in future_map:
                        pending.cancel()
                    self.status.emit("Операция отменена пользователем")
                    break
                self.status.emit(f"Конвертация: {file.name}")
                try:
                    converted_item = self._converted_file_item(future.result())
                    if converted_item is None:
                        raise Exception("pdf2docx не создал выходной DOCX-файл")
                    results.append(converted_item)
                except Exception as error:
                    self._report_file_conversion_error(file, error)
                finally:
                    processed += 1
                    self.progress.emit(int(processed / total * 100))

        return results

    def _warmup_word(self):
        if self._word_warmup_done:
            return
        if os.name != "nt" or not HAS_WORD_TO_PDF:
            self._word_warmup_done = True
            return
        try:
            started = prewarm_word_background(
                status_callback=(
                    self.status.emit
                    if callable(getattr(self, "status", None))
                    else None
                ),
                log_callback=_debug_log,
            )
            if started:
                time.sleep(0.6)
        except Exception as e:
            _debug_log(f"Не удалось подготовить Microsoft Word: {e}")
        finally:
            self._word_warmup_done = True

    def _convert_word_to_pdf(
        self,
        file: FileItem,
        *,
        output_reference: FileItem | None = None,
    ) -> str:
        if not file.path.lower().endswith((".doc", ".docx")):
            return None
        if not HAS_WORD_TO_PDF:
            raise Exception(
                "Конвертация Word в PDF недоступна. "
                "Установите pywin32 и убедитесь, что Microsoft Word установлен и активирован."
            )

        pdf_path = self._conversion_output_path(
            file,
            ".pdf",
            reference_file=output_reference,
        )
        try:
            self._warmup_word()
            if self._should_cancel():
                raise Exception("Конвертация отменена пользователем")
            if not self._convert_word_to_pdf_hidden_com(file.path, pdf_path):
                raise Exception("Microsoft Word не создал выходной PDF-файл")
            if self._should_cancel():
                raise Exception("Конвертация отменена пользователем")
            return pdf_path
        except Exception as error:
            err_text = str(error)
            if "could not start Microsoft Word" in err_text:
                err_text = (
                    "Не удалось запустить Microsoft Word. Убедитесь, что Word "
                    "установлен и активирован, хотя бы один раз запускался вручную "
                    "и не осталось зависших процессов WINWORD.EXE."
                )
            raise Exception(f"Ошибка конвертации Word в PDF: {err_text}") from error

    def _convert_word_to_pdf_hidden_com(self, src_path: str, dst_pdf_path: str) -> bool:
        if os.name != "nt":
            return False
        try:
            import pythoncom
            import win32com.client as win32
        except Exception:
            return False

        word_app = None
        document = None
        normalized_src = self._normalize_word_com_path(src_path)
        normalized_dst = self._normalize_word_com_path(dst_pdf_path)
        if not os.path.exists(normalized_src):
            raise FileNotFoundError(f"Файл не найден: {normalized_src}")
        pythoncom.CoInitialize()
        try:
            word_app = win32.DispatchEx("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = 0

            document = word_app.Documents.Open(
                normalized_src,
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                Revert=False,
                Visible=False,
                OpenAndRepair=True,
            )
            document.ExportAsFixedFormat(
                OutputFileName=normalized_dst,
                ExportFormat=_WD_EXPORT_FORMAT_PDF,
                OpenAfterExport=False,
            )
            return os.path.exists(normalized_dst)
        finally:
            _close_word_document(document)
            _quit_word_application(word_app)
            _uninitialize_com(pythoncom)

    def _convert_pdf_to_word(self, file: FileItem) -> str:
        if file.path.lower().endswith(".pdf"):
            docx_path = self._conversion_output_path(file, ".docx")

            if HAS_PDF_TO_WORD:
                try:
                    converter = pdf2docx.Converter(file.path)
                    converter.convert(docx_path)
                    converter.close()
                    return docx_path
                except Exception as e:
                    raise Exception(f"Ошибка конвертации PDF в Word: {e}")
            else:
                raise Exception("Установите pdf2docx")

        return None

    def _convert_word_to_odt(
        self,
        file: FileItem,
        *,
        output_reference: FileItem | None = None,
    ) -> str:
        if file.path.lower().endswith((".doc", ".docx")):
            if not HAS_ODF_PYTHON:
                raise Exception("Установите odfpy для конвертации ODT")
            try:
                import docx

                doc = docx.Document(file.path)
                odt_doc = OpenDocumentText()
                for para in doc.paragraphs:
                    p = text.P(text=para.text)
                    odt_doc.text.addElement(p)
                odt_path = self._conversion_output_path(
                    file,
                    ".odt",
                    reference_file=output_reference,
                )
                odt_doc.save(odt_path)
                return odt_path
            except Exception as e:
                raise Exception(f"Ошибка конвертации Word в ODT: {e}")
        return None

    def _convert_odt_to_word(self, file: FileItem) -> str:
        if file.path.lower().endswith(".odt"):
            if not HAS_ODF_PYTHON:
                raise Exception("Установите odfpy для конвертации ODT")
            try:
                import docx

                odt_doc = load(file.path)
                docx_doc = docx.Document()
                for elem in odt_doc.getElementsByType(text.P):
                    txt = teletype.extractText(elem)
                    if txt and txt.strip():
                        docx_doc.add_paragraph(txt)
                docx_path = self._conversion_output_path(file, ".docx")
                docx_doc.save(docx_path)
                return docx_path
            except Exception as e:
                raise Exception(f"Ошибка конвертации ODT в DOCX: {e}")
        return None

    def _convert_odt_to_pdf(self, file: FileItem) -> str:
        if file.path.lower().endswith(".odt"):
            docx_path = self._convert_odt_to_word(file)
            if not docx_path or not os.path.exists(docx_path):
                raise Exception("Не удалось создать DOCX из ODT")
            try:
                return self._convert_word_to_pdf(
                    FileItem(docx_path),
                    output_reference=file,
                )
            finally:
                try:
                    if os.path.exists(docx_path):
                        os.remove(docx_path)
                except OSError as error:
                    _debug_log(f"Не удалось удалить промежуточный DOCX: {error}")
        return None

    def _convert_pdf_to_odt(self, file: FileItem) -> str:
        if file.path.lower().endswith(".pdf"):
            docx_path = self._convert_pdf_to_word(file)
            if not docx_path or not os.path.exists(docx_path):
                raise Exception("Не удалось создать DOCX из PDF")
            try:
                return self._convert_word_to_odt(
                    FileItem(docx_path),
                    output_reference=file,
                )
            finally:
                try:
                    if os.path.exists(docx_path):
                        os.remove(docx_path)
                except OSError as error:
                    _debug_log(f"Не удалось удалить промежуточный DOCX: {error}")
        return None

    def _convert_pdf_to_image(self, file: FileItem) -> str:
        if not file.path.lower().endswith(".pdf"):
            return None
        if not HAS_PYMUPDF:
            raise Exception("Установите PyMuPDF для конвертации PDF в изображение")

        image_path = self._conversion_output_path(file, ".jpg")
        try:
            import pymupdf as fitz

            with fitz.open(file.path) as pdf_document:
                if pdf_document.page_count < 1:
                    raise Exception("PDF не содержит страниц")
                page = pdf_document.load_page(0)
                # ~200 DPI при стандартных 72 DPI PDF.
                pix = page.get_pixmap(matrix=fitz.Matrix(200 / 72, 200 / 72), alpha=False)
                pix.save(image_path)
            return image_path
        except Exception as error:
            raise Exception(f"Ошибка конвертации PDF в изображение: {error}") from error

    def _convert_image_to_pdf(self, file: FileItem) -> str:
        source_format = format_for_path(file.path)
        if source_format not in source_formats_for_category(IMAGE_CATEGORY):
            return None
        if not HAS_PIL:
            raise Exception("Установите Pillow для конвертации изображения в PDF")

        pdf_path = self._conversion_output_path(file, ".pdf")
        image = None
        try:
            if source_format == "SVG":
                image = self._load_svg_as_pillow_image(file.path)
            else:
                image = Image.open(file.path)
            frame_count = int(getattr(image, "n_frames", 1) or 1)
            if frame_count > 1:
                from PIL import ImageSequence

                frames = [
                    self._flatten_transparency(frame.copy())
                    for frame in ImageSequence.Iterator(image)
                ]
                first, rest = frames[0], frames[1:]
                first.save(pdf_path, "PDF", resolution=100.0, save_all=True, append_images=rest)
            else:
                frame = self._flatten_transparency(image)
                frame.save(pdf_path, "PDF", resolution=100.0)
            return pdf_path
        except Exception as error:
            self._discard_conversion_output(pdf_path)
            raise Exception(f"Ошибка конвертации изображения в PDF: {error}") from error
        finally:
            if image is not None:
                try:
                    image.close()
                except Exception as close_error:
                    _debug_log(f"Не удалось закрыть изображение {file.path}: {close_error}")

